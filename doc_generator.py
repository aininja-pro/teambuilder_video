import os
import tempfile
from datetime import datetime
from typing import List, Dict
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import streamlit as st
from parse_scope import TEAMBUILDERS_COST_CODES

# Try to import reportlab as fallback for PDF generation
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

def generate_docx(scope_items: List[Dict[str, str]], project_summary: Dict, job_name: str = "Job", version: int = 1) -> str:
    """
    Generate a DOCX document from scope items using TeamBuilders cost codes.
    
    Args:
        scope_items: List of formatted scope items
        project_summary: Dictionary containing project summary information
        job_name: Name of the job for the document title
        version: Version number for the document
        
    Returns:
        str: Path to the generated DOCX file
        
    Raises:
        Exception: If document generation fails
    """
    try:
        doc = Document()
        styles = doc.styles

        # --- FONT SIZE ADJUSTMENTS ---
        # Decrease font size for Normal text
        styles['Normal'].font.size = Pt(9)
        
        # Decrease font size for Headings
        styles['Heading 1'].font.size = Pt(14)
        styles['Heading 2'].font.size = Pt(12)
        styles['Title'].font.size = Pt(24) # Assuming a title style exists

        # --- HEADER ---
        title = doc.add_heading(f'Scope Summary: {job_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Reduce spacing after the title
        title.paragraph_format.space_after = Pt(12)
        
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        version_p = doc.add_paragraph(f'Version: {version}')
        
        # Reduce spacing after the version paragraph
        version_p.paragraph_format.space_after = Pt(18)
        
        # --- PROJECT SUMMARY SECTION ---
        summary_heading = doc.add_heading('Project Summary', level=1)
        summary_heading.paragraph_format.space_before = Pt(0) # Remove space above summary
        
        # Overview
        doc.add_heading('Overall Sentiment', level=2)
        doc.add_paragraph(project_summary.get('sentiment', 'N/A'))

        doc.add_heading('Overview', level=2)
        doc.add_paragraph(project_summary.get('overview', 'No overview provided.'))

        # Key Requirements
        doc.add_heading('Key Requirements', level=2)
        for item in project_summary.get('keyRequirements', []):
            doc.add_paragraph(item, style='List Bullet')

        # Concerns
        doc.add_heading('Concerns', level=2)
        for item in project_summary.get('concerns', []):
            doc.add_paragraph(item, style='List Bullet')

        # Decision Points
        doc.add_heading('Decision Points', level=2)
        for item in project_summary.get('decisionPoints', []):
            doc.add_paragraph(item, style='List Bullet')

        # Important Notes
        doc.add_heading('Important Notes', level=2)
        for item in project_summary.get('importantNotes', []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_page_break()

        # --- SCOPE ITEMS SECTION ---
        doc.add_heading('Detailed Scope Items', level=1)
        
        # --- GROUP SCOPE ITEMS BY MAIN CATEGORY ---
        grouped_items = {}
        for item in scope_items:
            main_code = item.get('Main Code', '00')
            main_category = item.get('Main Category', 'Uncategorized')
            
            group_key = f"{main_code} {main_category}"
            if group_key not in grouped_items:
                grouped_items[group_key] = []
            grouped_items[group_key].append(item)
            
        # Sort main categories by code
        sorted_groups = sorted(grouped_items.items(), key=lambda x: x[0])
        
        # --- ADD SCOPE ITEMS TO DOCUMENT ---
        for group_key, items in sorted_groups:
            # Add main category heading
            doc.add_heading(group_key, level=2)
            
            # Add sub-items
            for item in items:
                # Add sub-category and description
                sub_category = item.get('Sub Category', 'General')
                description = item.get('Description', 'No description provided.')
                
                p = doc.add_paragraph()
                p.add_run(f"{item.get('Sub Code')} {sub_category}: ").bold = True
                p.add_run(description)
                
                # Add details in a list or table for clarity
                details_to_add = {
                    "Material": item.get('Material'),
                    "Location": item.get('Location'),
                    "Quantity": item.get('Quantity'),
                    "Notes": item.get('Notes')
                }
                
                for key, value in details_to_add.items():
                    if value and value.strip(): # Only add if value exists
                        p_detail = doc.add_paragraph(style='List Bullet')
                        p_detail.add_run(f"{key}: ").bold = True
                        p_detail.add_run(f"{value}")

                doc.add_paragraph('') # Add space after each item
        
        # --- FOOTER ---
        doc.add_paragraph('')
        footer_para = doc.add_paragraph('--- End of Scope Summary ---')
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # --- SAVE DOCUMENT ---
        temp_dir = tempfile.gettempdir()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_filename = f"scope_summary_{job_name}_{timestamp}_v{version}.docx"
        docx_path = os.path.join(temp_dir, docx_filename)
        
        doc.save(docx_path)
        
        return docx_path
        
    except Exception as e:
        # Log the full traceback for debugging
        import traceback
        st.error(f"Error in DOCX generation: {e}")
        st.error(traceback.format_exc())
        raise Exception(f"DOCX generation failed: {str(e)}")

def generate_pdf(docx_path: str, job_name: str = "Job", version: int = 1) -> str:
    """
    Generate a PDF document from the DOCX file or create directly from scope items.
    
    Args:
        docx_path: Path to the DOCX file (for reference)
        job_name: Name of the job for the document title
        version: Version number for the document
        
    Returns:
        str: Path to the generated PDF file
        
    Raises:
        Exception: If PDF generation fails
    """
    try:
        if not REPORTLAB_AVAILABLE:
            raise Exception("PDF generation requires reportlab. Install with: pip install reportlab")
        
        # Create PDF path
        temp_dir = tempfile.gettempdir()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_filename = f"scope_summary_{job_name}_{timestamp}_v{version}.pdf"
        pdf_path = os.path.join(temp_dir, pdf_filename)
        
        # Create PDF document
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        story.append(Paragraph(f'Scope Summary: {job_name}', title_style))
        story.append(Spacer(1, 12))
        
        # Metadata
        story.append(Paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}', styles['Normal']))
        story.append(Paragraph(f'Version: {version}', styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Introduction
        story.append(Paragraph('Project Scope Overview', styles['Heading2']))
        story.append(Paragraph('This document contains the scope items extracted from the job site video, organized by construction division codes.', styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Note: For a complete PDF, we would need to re-parse the scope items
        # For now, we'll create a simple PDF that references the DOCX
        story.append(Paragraph('Scope Items', styles['Heading2']))
        story.append(Paragraph('Please refer to the DOCX file for detailed scope items, or contact support for a fully formatted PDF version.', styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Footer
        story.append(Spacer(1, 50))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            alignment=1,  # Center alignment
            fontSize=10
        )
        story.append(Paragraph('--- End of Scope Summary ---', footer_style))
        
        # Build PDF
        doc.build(story)
        
        return pdf_path
        
    except Exception as e:
        raise Exception(f"PDF generation failed: {str(e)}")

def generate_pdf_from_scope_items(scope_items: List[Dict[str, str]], project_summary: Dict, job_name: str = "Job", version: int = 1) -> str:
    """
    Generate a PDF document directly from scope items using TeamBuilders cost codes.
    
    Args:
        scope_items: List of formatted scope items from the parser
        project_summary: Dictionary containing project summary information
        job_name: Name of the job for the document title
        version: Version number for the document
        
    Returns:
        str: Path to the generated PDF file
    """
    try:
        if not REPORTLAB_AVAILABLE:
            raise Exception("PDF generation requires reportlab. Install with: pip install reportlab")
        
        # --- SETUP PDF ---
        temp_dir = tempfile.gettempdir()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_filename = f"scope_summary_{job_name}_{timestamp}_v{version}.pdf"
        pdf_path = os.path.join(temp_dir, pdf_filename)
        
        doc = SimpleDocTemplate(pdf_path, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []

        # --- STYLES (with adjusted font sizes) ---
        title_style = ParagraphStyle('CustomTitle', parent=styles['h1'], fontSize=16, alignment=1, spaceAfter=18)
        h1_style = ParagraphStyle('CustomH1', parent=styles['h1'], fontSize=14, spaceBefore=12, spaceAfter=8)
        h2_style = ParagraphStyle('CustomH2', parent=styles['h2'], fontSize=12, spaceBefore=10, spaceAfter=6)
        body_style = ParagraphStyle('CustomBody', parent=styles['BodyText'], fontSize=8)
        bullet_style = ParagraphStyle('CustomBullet', parent=body_style, leftIndent=20, spaceAfter=4)

        # --- HEADER ---
        story.append(Paragraph(f'Scope Summary: {job_name}', title_style))
        story.append(Paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}', body_style))
        story.append(Paragraph(f'Version: {version}', body_style))
        story.append(Spacer(1, 0.2*inch))
        
        # --- PROJECT SUMMARY SECTION ---
        story.append(Paragraph('Project Summary', h1_style))
        
        story.append(Paragraph(f"<b>Overall Sentiment:</b> {project_summary.get('sentiment', 'N/A')}", body_style))
        story.append(Spacer(1, 0.1*inch))
        
        story.append(Paragraph('Overview', h2_style))
        story.append(Paragraph(project_summary.get('overview', 'No overview provided.'), body_style))
        story.append(Spacer(1, 0.2*inch))
        
        def add_list_items(title, items):
            if items:
                story.append(Paragraph(title, h2_style))
                for item in items:
                    story.append(Paragraph(f"• {item}", bullet_style))
                story.append(Spacer(1, 0.2*inch))

        add_list_items('Key Requirements', project_summary.get('keyRequirements', []))
        add_list_items('Concerns', project_summary.get('concerns', []))
        add_list_items('Decision Points', project_summary.get('decisionPoints', []))
        add_list_items('Important Notes', project_summary.get('importantNotes', []))

        story.append(PageBreak())

        # --- SCOPE ITEMS SECTION ---
        story.append(Paragraph('Detailed Scope Items', h1_style))
        story.append(Spacer(1, 0.2*inch))

        # --- GROUP AND ADD SCOPE ITEMS ---
        grouped_items = {}
        for item in scope_items:
            main_code = item.get('Main Code', '00')
            main_category = item.get('Main Category', 'Uncategorized')
            group_key = f"{main_code} {main_category}"
            if group_key not in grouped_items:
                grouped_items[group_key] = []
            grouped_items[group_key].append(item)

        sorted_groups = sorted(grouped_items.items(), key=lambda x: x[0])

        for group_key, items in sorted_groups:
            story.append(Paragraph(group_key, h2_style))
            for item in items:
                sub_category = item.get('Sub Category', 'General')
                description = item.get('Description', 'No description provided.')
                
                story.append(Paragraph(f"<b>{item.get('Sub Code')} {sub_category}:</b> {description}", body_style))
                
                details_to_add = {
                    "Material": item.get('Material'),
                    "Location": item.get('Location'),
                    "Quantity": item.get('Quantity'),
                    "Notes": item.get('Notes')
                }
                
                for key, value in details_to_add.items():
                    if value and value.strip():
                        story.append(Paragraph(f"• <b>{key}:</b> {value}", bullet_style))
                story.append(Spacer(1, 0.1*inch))
        
        # --- FOOTER ---
        story.append(Spacer(1, 0.5*inch))
        footer_style = ParagraphStyle('Footer', parent=body_style, alignment=1)
        story.append(Paragraph('--- End of Scope Summary ---', footer_style))
        
        # --- BUILD DOCUMENT ---
        doc.build(story)
        return pdf_path

    except Exception as e:
        import traceback
        st.error(f"Error in PDF generation: {e}")
        st.error(traceback.format_exc())
        raise Exception(f"PDF generation failed: {str(e)}")

def create_filename(job_name: str, version: int, extension: str) -> str:
    """
    Create a standardized filename for generated documents.
    
    Args:
        job_name: Name of the job
        version: Version number
        extension: File extension (e.g., 'docx', 'pdf')
        
    Returns:
        str: Formatted filename
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_job_name = "".join(c for c in job_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
    safe_job_name = safe_job_name.replace(' ', '_')
    
    return f"{safe_job_name}_ScopeSummary_{timestamp}_v{version}.{extension}" 
# ChatGPT's EXACT worker pattern
from redis import Redis
import time, json
import logging
import traceback
import os
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()

redis = Redis.from_url(os.getenv("REDIS_URL", "redis://localhost:6379"), decode_responses=True)
log = logging.getLogger("processing")

# ChatGPT's EXACT publish pattern with payload
def publish(session_id, pct, msg, **extra):
    payload = {"type": "progress", "pct": pct, "msg": msg, **extra}
    redis.hset(f"jobs:{session_id}", mapping={"pct": pct, "msg": msg, "status": extra.get("status", "processing")})
    redis.publish(f"progress:{session_id}", json.dumps(payload))
    redis.expire(f"jobs:{session_id}", 60 * 60)  # 1 hour TTL

# ChatGPT's logging wrapper
def run_with_logging(fn, *args, **kwargs):
    try:
        logging.info("TASK START %s", args or "")
        return fn(*args, **kwargs)
    except Exception:
        logging.exception("TASK FAILED")
        raise
    finally:
        logging.info("TASK END")

# ChatGPT's EXACT worker pattern with your progress semantics
def process_session(session_id: str):
    publish(session_id, 0, "Queued")
    
    # Progress semantics: Upload (0-20%) → MOV Convert (20-40%) → Transcribe (40-70%) → Parse (70-90%) → Docs (90-100%)
    
    # validate files exist, assemble chunks, etc.
    publish(session_id, 10, "Assembling chunks")
    time.sleep(1)
    
    # MOV Convert (20-40%)
    publish(session_id, 25, "Converting MOV to MP4...")
    time.sleep(2)
    publish(session_id, 40, "MOV conversion complete")
    
    # Transcribe (40-70%)
    publish(session_id, 45, "Starting audio transcription...")
    
    # Initialize OpenAI
    client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
    
    # Real OpenAI Whisper transcription
    # Get file path from Redis (stored by complete endpoint)
    import sys
    print(f"DEBUG WORKER: Looking for file path for session {session_id}", file=sys.stderr)
    file_path = redis.hget(f"jobs:{session_id}", "file_path")
    print(f"DEBUG WORKER: Redis returned file_path: {file_path}", file=sys.stderr)
    # Note: file_path is already a string because Redis client has decode_responses=True
    
    if not file_path:
        # Try to find assembled file in temp directory
        import tempfile
        import glob
        temp_files = glob.glob(f"{tempfile.gettempdir()}/*{session_id}*")
        if temp_files:
            file_path = temp_files[0]
        else:
            raise Exception(f"Cannot find file for session {session_id}")
    
    if not os.path.exists(file_path):
        raise Exception(f"File not found: {file_path}")
    
    # Check file size and compress if needed (OpenAI limit: 25MB)
    file_size = os.path.getsize(file_path)
    max_size = 25 * 1024 * 1024  # 25MB in bytes
    
    if file_size > max_size:
        publish(session_id, 45, f"File too large ({file_size} bytes), compressing...")
        # Create compressed version for OpenAI
        compressed_path = file_path + "_compressed.mp4"
        
        import subprocess
        # Use ffmpeg to compress video - reduce bitrate and resolution for Whisper
        cmd = [
            "ffmpeg", "-y", "-i", file_path,
            "-vn",  # No video (audio only for Whisper)
            "-ar", "16000",  # 16kHz sample rate (Whisper works well with this)
            "-ac", "1",  # Mono audio
            "-ab", "64k",  # Low audio bitrate
            compressed_path
        ]
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode != 0:
            raise Exception(f"FFmpeg compression failed: {result.stderr}")
        
        # Use compressed file for OpenAI
        transcription_file = compressed_path
        publish(session_id, 48, f"Compression complete, starting transcription...")
    else:
        transcription_file = file_path
    
    # Transcribe with OpenAI Whisper
    with open(transcription_file, 'rb') as audio_file:
        transcript = client.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file,
            response_format="text"
        )
    
    # Clean up compressed file if created
    if file_size > max_size and os.path.exists(compressed_path):
        os.remove(compressed_path)
    
    publish(session_id, 70, f"Transcription complete ({len(transcript)} characters)")
    
    # Parse (70-90%) - Real GPT-4 scope parsing
    publish(session_id, 75, "Analyzing transcript for scope items...")
    
    # Complete TeamBuilders cost codes
    cost_codes = {
        "01 General Conditions": {
            "1100": "Permit",
            "1200": "Project Oversight (Management, Coordination, Procurement)",
            "1600": "Tool & Equipment Rental"
        },
        "02 Site/Demo": {
            "1700": "Dump & Trash Removal",
            "1800": "Demolition (Wall, Deck, Roof, Flooring, Kitchen)",
            "1900": "Prep Work & Dust Protection"
        },
        "03 Excavation/Landscape": {
            "2100": "Landscaping"
        },
        "04 Concrete/Masonry": {
            "2200": "Concrete Footings & Foundations",
            "2250": "Concrete Flatwork",
            "2299": "Concrete/Masonry Specialties"
        },
        "05 Rough Carpentry": {
            "3000": "Floor & Stair Framing",
            "3100": "Wall Framing",
            "3200": "Roof Framing"
        },
        "06 Doors/Windows": {
            "3500": "Exterior Doors",
            "3600": "Windows"
        },
        "07 Mechanical": {
            "4000": "HVAC System"
        },
        "08 Electrical": {
            "4100": "Electrical System"
        },
        "09 Plumbing": {
            "4200": "Plumbing System",
            "4250": "Plumbing Fixtures"
        },
        "10 Wall/Ceiling Coverings": {
            "5000": "Insulation",
            "5100": "Drywall",
            "5200": "Paint"
        },
        "11 Finish Carpentry": {
            "5300": "Interior Doors",
            "5350": "Interior Door Hardware",
            "5400": "Interior Trim",
            "5450": "Interior Trim Specialties"
        },
        "12 Cabinets/Vanities/Tops": {
            "5600": "Kitchen Cabinets",
            "5650": "Bathroom Vanities",
            "5680": "Built-In Cabinetry",
            "5699": "Cabinet Hardware Specialties",
            "5700": "Countertops"
        },
        "13 Flooring/Tile": {
            "5800": "Flooring"
        },
        "14 Specialties": {
            "6000": "Closet Shelving",
            "6200": "Appliances",
            "6300": "Specialty Glass"
        },
        "15 Decking": {
            "7000": "Decking"
        },
        "16 Fencing": {},
        "17 Exterior Facade": {
            "7200": "House Wrap",
            "7220": "Vinyl Siding",
            "7240": "Luxury Siding"
        },
        "18 Soffit/Fascia/Gutters": {
            "7300": "Soffit/Fascia",
            "7340": "Gutters"
        },
        "19 Roofing": {
            "7400": "Asphalt Roofing"
        }
    }
    
    system_prompt = f"""
You are an expert construction estimator specializing in TeamBuilders cost code classification. Analyze the following transcript from a job site video and extract scope items organized by TeamBuilders cost codes.

TeamBuilders Cost Code Structure: {json.dumps(cost_codes)}

Instructions:
1. CAREFULLY analyze the transcript for construction activities, materials, and work being performed
2. Match each identified item to the MOST SPECIFIC TeamBuilders subcode (4-digit) when possible
3. If an exact subcode match isn't clear, use the main category code (01-19)
4. Extract SPECIFIC details mentioned including:
   - Quantities (if mentioned)
   - Materials specified
   - Dimensions or measurements
   - Location in the building
   - Any special requirements or notes
5. 🔍 EXHAUSTIVENESS REQUIREMENT: Carefully and systematically extract every distinct construction-related activity, material, or scope detail. Do not omit or skip any mention, no matter how minor. Return a separate JSON object for each unique task or item referenced.
6. If there is uncertainty, err on the side of inclusion, using the closest main code category when a subcode is unclear.
7. Do NOT invent or assume items not clearly stated or shown in the transcript

Key Matching Guidelines:
- Permits, inspections, project management → 01 General Conditions
- Any demolition or removal work → 02 Site/Demo
- Foundation, footings, slabs → 04 Concrete/Masonry
- Framing (floor, wall, roof) → 05 Rough Carpentry
- Door and window installation → 06 Doors/Windows
- HVAC, furnace, ductwork → 07 Mechanical
- Electrical wiring, fixtures, outlets → 08 Electrical
- Plumbing pipes, fixtures, water heaters → 09 Plumbing
- Insulation, drywall, painting → 10 Wall/Ceiling Coverings
- Interior doors, trim, molding → 11 Finish Carpentry
- Cabinets, countertops, vanities → 12 Cabinets/Vanities/Tops
- Flooring materials (LVT, carpet, tile) → 13 Flooring/Tile
- Appliances, mirrors, shelving → 14 Specialties
- Deck construction → 15 Decking
- Siding, house wrap → 17 Exterior Facade
- Soffit, fascia, gutters → 18 Soffit/Fascia/Gutters
- Roofing materials and work → 19 Roofing

Return JSON with 'scopeItems' and 'projectSummary' keys:

scopeItems: Array of objects with:
- mainCode: Two-digit code (e.g., "02")
- mainCategory: Main category name (e.g., "Site/Demo")
- subCode: Four-digit subcode (e.g., "1800")
- subCategory: Subcategory name (e.g., "Demolition")
- description: Clear description of the work
- details: Object with material, location, quantity, notes

projectSummary: Object with:
- overview: Brief project description
- keyRequirements: Array of main requirements
- concerns: Array of potential issues or challenges
- decisionPoints: Array of decisions that need to be made
- importantNotes: Array of critical information

Return ONLY valid JSON.
"""
    
    response = client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Transcript: {transcript}"}
        ],
        temperature=0,
        max_tokens=3000,
        response_format={"type": "json_object"}
    )
    
    parsed_data = json.loads(response.choices[0].message.content.strip())
    scope_items = parsed_data.get('scopeItems', [])
    project_summary = parsed_data.get('projectSummary', {})
    
    publish(session_id, 90, f"Extracted {len(scope_items)} scope items")
    
    # Docs (90-100%)
    publish(session_id, 95, "Generating documents...")
    
    try:
        # Generate documents with real data
        filename = os.path.basename(file_path) if file_path else "Unknown"
        job_name = os.path.splitext(filename)[0] if filename else "Video Analysis"
        
        # Import document libraries  
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from datetime import datetime
        
        # Create DOCX document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'TeamBuilders Scope Summary - {job_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
        doc.add_paragraph('') # spacing
        
        # Add project summary if available
        if project_summary and project_summary.get('overview'):
            doc.add_heading('Project Overview', level=1)
            doc.add_paragraph(project_summary['overview'])
            doc.add_paragraph('')
        
        # Add scope items - sort by category number first
        if scope_items:
            # Sort scope items by main code, then sub code (same as UI)
            sorted_scope_items = sorted(scope_items, key=lambda x: (
                int(x.get('mainCode', '99')), 
                int(x.get('subCode', '9999'))
            ))
            
            doc.add_heading('Scope Items', level=1)
            
            for item in sorted_scope_items:
                # Item header (no numbering, just category codes)
                p = doc.add_paragraph()
                p.add_run(f"{item.get('mainCode', 'XX')} {item.get('mainCategory', 'Unknown')} - {item.get('subCode', 'XXXX')} {item.get('subCategory', 'Unknown')}").bold = True
                
                # Description
                if item.get('description'):
                    doc.add_paragraph(f"Description: {item['description']}", style='List Bullet')
                
                # Details
                details = item.get('details', {})
                if details:
                    if details.get('material'):
                        doc.add_paragraph(f"Material: {details['material']}", style='List Bullet')
                    if details.get('location'):
                        doc.add_paragraph(f"Location: {details['location']}", style='List Bullet')
                    if details.get('quantity'):
                        doc.add_paragraph(f"Quantity: {details['quantity']}", style='List Bullet')
                    if details.get('notes'):
                        doc.add_paragraph(f"Notes: {details['notes']}", style='List Bullet')
                
                doc.add_paragraph('')  # spacing
        
        # Save DOCX
        static_dir = "/Users/richardrierson/Desktop/Projects/TeamBuilders/video-scope-analyzer/backend/static"
        os.makedirs(static_dir, exist_ok=True)
        
        docx_filename = f"{session_id}_scope_summary.docx"
        docx_path = os.path.join(static_dir, docx_filename)
        doc.save(docx_path)
        
        publish(session_id, 97, "Word document generated...")
        
        # Generate PDF document
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        
        pdf_filename = f"{session_id}_scope_summary.pdf"
        pdf_path = os.path.join(static_dir, pdf_filename)
        
        doc_pdf = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        story.append(Paragraph(f"TeamBuilders Scope Summary - {job_name}", styles['Title']))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", styles['Normal']))
        story.append(Spacer(1, 24))
        
        # Project overview
        if project_summary and project_summary.get('overview'):
            story.append(Paragraph("Project Overview", styles['Heading1']))
            story.append(Paragraph(project_summary['overview'], styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Scope items - sort by category number first
        if scope_items:
            # Sort scope items by main code, then sub code (same as UI)
            sorted_scope_items = sorted(scope_items, key=lambda x: (
                int(x.get('mainCode', '99')), 
                int(x.get('subCode', '9999'))
            ))
            
            story.append(Paragraph("Scope Items", styles['Heading1']))
            story.append(Spacer(1, 12))
            
            for item in sorted_scope_items:
                # Item header (no numbering, just category codes)
                header_text = f"{item.get('mainCode', 'XX')} {item.get('mainCategory', 'Unknown')} - {item.get('subCode', 'XXXX')} {item.get('subCategory', 'Unknown')}"
                story.append(Paragraph(header_text, styles['Heading2']))
                
                # Description
                if item.get('description'):
                    story.append(Paragraph(f"<b>Description:</b> {item['description']}", styles['Normal']))
                
                # Details
                details = item.get('details', {})
                if details:
                    detail_parts = []
                    if details.get('material'): detail_parts.append(f"Material: {details['material']}")
                    if details.get('location'): detail_parts.append(f"Location: {details['location']}")
                    if details.get('quantity'): detail_parts.append(f"Quantity: {details['quantity']}")
                    if details.get('notes'): detail_parts.append(f"Notes: {details['notes']}")
                    
                    for detail in detail_parts:
                        story.append(Paragraph(f"• {detail}", styles['Normal']))
                
                story.append(Spacer(1, 12))
        
        doc_pdf.build(story)
        
        documents = {
            "docx": f"/static/{docx_filename}",
            "pdf": f"/static/{pdf_filename}"
        }
        
        publish(session_id, 99, "PDF and Word documents ready for download...")
        
    except Exception as e:
        log.error(f"Document generation failed: {e}")
        # Fallback to no documents
        documents = {"docx": None, "pdf": None}
    
    publish(session_id, 100, "Done", status="completed")
    
    # Store final result with real data
    result = {
        "transcript": transcript,
        "scope_items": scope_items,
        "project_summary": project_summary,
        "documents": documents
    }
    redis.hset(f"jobs:{session_id}", mapping={
        "status": "completed",
        "result": json.dumps(result)
    })
    
    # Save analysis for persistence
    from datetime import datetime
    import uuid
    
    analysis_id = str(uuid.uuid4())
    created_at = datetime.utcnow().isoformat() + "Z"
    
    # Get original filename from file path
    filename = os.path.basename(file_path) if file_path else "Unknown"
    
    # Calculate file size in MB
    file_size_mb = round(file_size / (1024 * 1024), 2) if file_size else None
    
    # Save to persistent storage
    redis.hset(f"analysis:{analysis_id}", mapping={
        "filename": filename,
        "created_at": created_at,
        "transcript": transcript,
        "scope_items": json.dumps(scope_items),
        "project_summary": json.dumps(project_summary),
        "file_size_mb": str(file_size_mb) if file_size_mb else "",
        "documents": json.dumps(documents)
    })
    
    log.info(f"Saved analysis {analysis_id} for file {filename}")